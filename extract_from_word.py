from docx import Document

def is_highlighted(run):
    """
    Überprüfe, ob der Text in einem Run hervorgehoben ist und gebe die Hervorhebungsfarbe aus.
    """
    if run.font.highlight_color is not None:
        print(f"Gefundene Hervorhebungsfarbe: {run.font.highlight_color}")  # Debug-Ausgabe
        return True

    return False

def is_red(run):
    """
    Überprüfe, ob der Text in einem Run rot eingefärbt ist.
    """
    if run.font.color is not None:
        # Farbe wird als RGB-Wert gespeichert, Rot ist (255, 0, 0)
        return run.font.color.rgb == (255, 0, 0)
    return False

def extract_highlighted_words_from_paragraph(paragraph):
    """
    Extrahiere hervorgehobene Wörter aus einem Absatz.

    :param paragraph: Ein Word-Absatz
    :return: Liste der hervorgehobenen Wörter im Absatz
    """
    highlighted_words = []
    # Durch alle Runs im Absatz gehen
    for run in paragraph.runs:
        if is_highlighted(run) or is_red(run):
            highlighted_words.append(run.text)
    return highlighted_words

def extract_highlighted_words(docx_path):
    """
    Extrahiere Wörter, die in einer .docx-Datei hervorgehoben sind.

    :param docx_path: Pfad zur .docx-Datei
    :return: Liste der hervorgehobenen Wörter
    """
    # Öffne das Dokument
    doc = Document(docx_path)
    
    highlighted_words = []

    # Durch alle Absätze im Dokument gehen
    for para in doc.paragraphs:
        highlighted_words.extend(extract_highlighted_words_from_paragraph(para))

    # Durch alle Tabellen im Dokument gehen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Auch hier die hervorgehobenen Wörter aus den Zellen extrahieren
                for para in cell.paragraphs:
                    highlighted_words.extend(extract_highlighted_words_from_paragraph(para))

    return highlighted_words
"""
# Beispiel-Aufruf des Skripts
if __name__ == "__main__":
    docx_file = "21_10_24_Halloween_Spukhaus_Plinius.docx"  # Pfad zur .docx-Datei
    highlighted_words = extract_highlighted_words(docx_file)
    
    # Ausgabe der hervorgehobenen Wörter
    print("Hervorgehobene Wörter im Dokument:", highlighted_words)"""